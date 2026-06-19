#!/usr/bin/env python3
"""生成 Pandoc 引用文档：纯黑文字、表格带边框、紧凑排版适合打印"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(SCRIPT_DIR, "..", "docx", "_reference.docx")


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    """设置段落间距和行距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def set_style_spacing(style, before=0, after=0, line_spacing=1.0):
    """设置样式级别的间距"""
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def set_page_margins(doc, top=1.5, bottom=1.5, left=1.5, right=1.5):
    """设置页面边距（单位：厘米）"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def set_table_borders(table):
    """为表格设置完整黑色边框"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))


def create_reference():
    doc = Document()

    # ─── 页面边距：1.5cm ───
    set_page_margins(doc, top=1.5, bottom=1.5, left=1.5, right=1.5)

    # ─── 全局样式 ───
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Noto Sans SC"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0, 0, 0)
    font.bold = False
    set_style_spacing(style, before=0, after=2, line_spacing=1.15)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Noto Sans SC")

    # ─── 标题样式 ───
    for i in range(1, 4):
        heading = doc.styles[f"Heading {i}"]
        hfont = heading.font
        hfont.color.rgb = RGBColor(0, 0, 0)
        hfont.bold = True
        hfont.name = "Noto Sans SC"
        hrPr = heading.element.get_or_add_rPr()
        hrFonts = hrPr.find(qn("w:rFonts"))
        if hrFonts is None:
            hrFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            hrPr.insert(0, hrFonts)
        hrFonts.set(qn("w:eastAsia"), "Noto Sans SC")

    doc.styles["Heading 1"].font.size = Pt(18)
    set_style_spacing(doc.styles["Heading 1"], before=8, after=4, line_spacing=1.15)
    doc.styles["Heading 2"].font.size = Pt(14)
    set_style_spacing(doc.styles["Heading 2"], before=6, after=3, line_spacing=1.15)
    doc.styles["Heading 3"].font.size = Pt(12)
    set_style_spacing(doc.styles["Heading 3"], before=4, after=2, line_spacing=1.15)

    # ─── 引用块样式 ───
    quote = doc.styles["Quote"]
    qfont = quote.font
    qfont.color.rgb = RGBColor(0, 0, 0)
    qfont.italic = False
    qfont.size = Pt(10)
    set_style_spacing(quote, before=2, after=2, line_spacing=1.15)

    # ─── 正文示例 ───
    p = doc.add_paragraph("正文段落示例。")
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.15)

    # ─── 标题示例 ───
    doc.add_heading("一级标题示例", level=1)
    doc.add_heading("二级标题示例", level=2)
    doc.add_heading("三级标题示例", level=3)

    # ─── 引用块示例 ───
    doc.add_paragraph("引用块示例文字。", style="Quote")

    # ─── 列表示例 ───
    doc.add_paragraph("列表项一", style="List Bullet")
    doc.add_paragraph("列表项二", style="List Bullet")

    # ─── 表格示例 ───
    table = doc.add_table(rows=3, cols=3)
    headers = ["表头一", "表头二", "表头三"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            set_paragraph_spacing(paragraph, before=1, after=1, line_spacing=1.0)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(10)

    data = [["内容A1", "内容B1", "内容C1"], ["内容A2", "内容B2", "内容C2"]]
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=1, after=1, line_spacing=1.0)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.size = Pt(10)

    set_table_borders(table)

    # ─── 代码块示例 ───
    cp = doc.add_paragraph()
    cp.text = "代码块示例"
    set_paragraph_spacing(cp, before=2, after=2, line_spacing=1.0)
    for run in cp.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"  引用文档生成: {OUT}")


if __name__ == "__main__":
    create_reference()
